"""Carga de archivos locales (CSV/Excel/PDF/TXT/MD/DOCX) -> DataFrame o texto + metadatos."""

from __future__ import annotations

from dataclasses import dataclass
from pathlib import Path

import pandas as pd

from app.schemas.data_schema import FileMetadata
from app.utils.file_utils import classify_document_type, file_extension, file_kind
from app.utils.logging import get_logger

logger = get_logger(__name__)


@dataclass
class LoadedTable:
    """Resultado de cargar un archivo tabular."""

    name: str
    df: pd.DataFrame
    metadata: FileMetadata


@dataclass
class LoadedDocument:
    """Resultado de cargar un documento textual."""

    name: str
    text: str
    metadata: FileMetadata


def load_table(path: str | Path) -> LoadedTable:
    """Carga un CSV/Excel a un DataFrame de pandas."""
    path = Path(path)
    ext = file_extension(path)
    if ext == ".csv":
        df = _read_csv_robust(path)
    elif ext in (".xlsx", ".xls"):
        df = pd.read_excel(path)
    else:
        raise ValueError(f"Extensión no tabular: {ext}")

    name = path.stem
    meta = FileMetadata(
        file_name=path.name,
        file_type=ext.lstrip("."),
        kind="tabular",
        rows=int(df.shape[0]),
        columns=int(df.shape[1]),
        status="loaded",
    )
    logger.info("Tabla cargada: %s (%d filas, %d cols)", path.name, df.shape[0], df.shape[1])
    return LoadedTable(name=name, df=df, metadata=meta)


def _read_csv_robust(path: Path) -> pd.DataFrame:
    """Lee un CSV probando codificaciones y separadores comunes."""
    for enc in ("utf-8", "utf-8-sig", "latin-1"):
        try:
            return pd.read_csv(path, encoding=enc, sep=None, engine="python")
        except (UnicodeDecodeError, pd.errors.ParserError):
            continue
    # Último intento simple
    return pd.read_csv(path)


def load_dataframe_from_bytes(content: bytes, file_name: str) -> LoadedTable:
    """Carga un DataFrame desde bytes (subida vía API)."""
    import io

    ext = file_extension(file_name)
    buffer = io.BytesIO(content)
    if ext == ".csv":
        try:
            df = pd.read_csv(buffer, sep=None, engine="python")
        except (UnicodeDecodeError, pd.errors.ParserError):
            buffer.seek(0)
            df = pd.read_csv(io.BytesIO(content), encoding="latin-1", sep=None, engine="python")
    elif ext in (".xlsx", ".xls"):
        df = pd.read_excel(buffer)
    else:
        raise ValueError(f"Extensión no tabular: {ext}")

    meta = FileMetadata(
        file_name=file_name,
        file_type=ext.lstrip("."),
        kind="tabular",
        rows=int(df.shape[0]),
        columns=int(df.shape[1]),
    )
    return LoadedTable(name=Path(file_name).stem, df=df, metadata=meta)


def load_document(path: str | Path) -> LoadedDocument:
    """Extrae texto de un documento (PDF/TXT/MD/DOCX)."""
    path = Path(path)
    ext = file_extension(path)
    text = extract_text(path.read_bytes(), path.name)
    meta = FileMetadata(
        file_name=path.name,
        file_type=ext.lstrip("."),
        kind="document",
        doc_type=classify_document_type(path.name),
        status="loaded",
    )
    return LoadedDocument(name=path.stem, text=text, metadata=meta)


def extract_text(content: bytes, file_name: str) -> str:
    """Extrae texto plano desde bytes según la extensión."""
    import io

    ext = file_extension(file_name)
    if ext in (".txt", ".md"):
        for enc in ("utf-8", "utf-8-sig", "latin-1"):
            try:
                return content.decode(enc)
            except UnicodeDecodeError:
                continue
        return content.decode("utf-8", errors="ignore")
    if ext == ".pdf":
        from pypdf import PdfReader

        reader = PdfReader(io.BytesIO(content))
        return "\n".join((page.extract_text() or "") for page in reader.pages)
    if ext == ".docx":
        import docx

        document = docx.Document(io.BytesIO(content))
        return "\n".join(p.text for p in document.paragraphs)
    raise ValueError(f"Extensión documental no soportada: {ext}")


def detect_kind(file_name: str) -> str:
    """Devuelve 'tabular', 'document' o 'unknown'."""
    return file_kind(file_name)
